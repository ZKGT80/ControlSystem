import os
import math
import wikipedia
from datetime import datetime
from docx import Document
from langchain_openai import ChatOpenAI

# 1. Connect to your running LM Studio server
llm = ChatOpenAI(
    base_url="http://127.0.0.1:1234/v1",
    api_key="lm-studio",
    model="hermes-2-pro-llama-3-8b",     
    temperature=0.2
)

# 2. Define the Word Document Creator function
def create_word_doc(content, filename="generated.docx", folder="generated_docs"):
    os.makedirs(folder, exist_ok=True)
    filepath = os.path.join(folder, filename)
    doc = Document()
    doc.add_heading("Generated Document", 0)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filepath)
    return f"Document saved at: {filepath}"

# 3. Core Tool Functions
def word_tool(query):
    prompt = f"Create a well-structured document:\n\n{query}\n\nFormat:\n- Title\n- Sections\n- Bullet points if needed"
    content = llm.invoke(prompt).content
    filename = f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return create_word_doc(content, filename)

def tool_calculator(query):
    try:
        # Sanitizing simple calculation input
        sanitized = "".join(c for c in query if c in "0123456789+-*/(). ")
        result = eval(sanitized)
        return str(result)
    except:
        return "Invalid arithmetic expression"

def wiki_doc_creator(query):
    print(f"-> Searching Wikipedia for '{query}'...")
    try:
        summary = wikipedia.summary(query, sentences=10)
    except Exception as e:
        return f"Could not extract Wikipedia data: {str(e)}"
        
    print("-> Formatting and building document...")
    prompt = f"Create a beautiful, deeply structured educational document using this background information:\n\n{summary}\n\nFormat:\n- Clear Heading Title\n- Comprehensive Subsections\n- Key Bullet points and summaries where needed"
    content = llm.invoke(prompt).content
    filename = f"wiki_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return create_word_doc(content, filename)

# 4. Light Routing Logic (Acts exactly like your classmate's agent decision rules)
def run_agent(user_input):
    decision_prompt = f"""Analyze this user request: "{user_input}"
Determine which task type it belongs to. Respond with ONLY the number corresponding to the choice:
1 - The user wants to pull data from Wikipedia and create a Word Document out of it.
2 - The user just wants to write a general custom document without using Wikipedia.
3 - The user is asking for an arithmetic calculation.
4 - The user is asking a basic general knowledge question.

Response (1, 2, 3, or 4):"""
    
    try:
        choice = llm.invoke(decision_prompt).content.strip()
        
        if "1" in choice:
            # Clean up query to extract just the topic name
            cleanup = llm.invoke(f"Extract only the main topic subject from this request as a search term: '{user_input}'. Respond with only the topic word(s).").content.strip()
            return wiki_doc_creator(cleanup)
        elif "2" in choice:
            return word_tool(user_input)
        elif "3" in choice:
            calc_input = llm.invoke(f"Extract just the math equation from this text: '{user_input}'. Example: 4+4. Output nothing else.").content.strip()
            return f"Calculation Result: {tool_calculator(calc_input)}"
        else:
            return llm.invoke(user_input).content
    except Exception as e:
        return f"Routing Error: {str(e)}"

# 5. Interactive Terminal Loop
if __name__ == "__main__":
    print("="*60)
    print("Classmate's Automated Wiki-Doc System Active!")
    print("="*60)
    
    while True:
        query = input("Ask a question (type 'exit' to quit): ")
        if query.lower() == "exit":
            print("Exiting... Goodbye!")
            break
        if not query.strip():
            continue
            
        print("\n[Agent Processing Request...]")
        response = run_agent(query)
        print("-" * 60)
        print("AI Response:\n>>", response)
        print("-" * 60 + "\n")